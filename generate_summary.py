"""
Generate Intelli-Credit One-Page Summary Document
FinTech-2: Loan Default Risk Prediction System
National Level AI Grand Challenge 2026
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# ── Colors ──
TEAL = RGBColor(0x0D, 0x94, 0x88)
DARK_NAVY = RGBColor(0x0F, 0x17, 0x2A)
SLATE_700 = RGBColor(0x33, 0x41, 0x55)
SLATE_500 = RGBColor(0x64, 0x74, 0x8B)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_TEAL_BG = "D1FAE5"  # emerald-100
TEAL_BG = "CCFBF1"         # teal-100
AMBER_BG = "FEF3C7"        # amber-100
RED_BG = "FEE2E2"          # red-100
BLUE_BG = "DBEAFE"         # blue-100
PURPLE_BG = "EDE9FE"       # violet-100
GRAY_BG = "F1F5F9"         # slate-100
NAVY_BG = "1E293B"         # slate-800


def set_cell_bg(cell, color_hex):
    """Set cell background color."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_border(cell, **kwargs):
    """Set cell borders."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, val in kwargs.items():
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="{val.get("val", "single")}" '
            f'w:sz="{val.get("sz", "4")}" w:space="0" '
            f'w:color="{val.get("color", "000000")}"/>'
        )
        tcBorders.append(element)
    tcPr.append(tcBorders)


def add_styled_paragraph(doc_or_cell, text, font_size=9, bold=False, color=SLATE_700,
                          alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=2, space_before=0,
                          font_name="Calibri"):
    """Add a styled paragraph."""
    p = doc_or_cell.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing = Pt(font_size + 2)
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = font_name
    return p


def add_bullet(doc_or_cell, text, font_size=8, color=SLATE_700, bold_prefix="", indent=0.25):
    """Add a bullet point."""
    p = doc_or_cell.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = Pt(font_size + 2.5)
    p.paragraph_format.left_indent = Inches(indent)
    p.paragraph_format.first_line_indent = Inches(-0.15)
    
    run_bullet = p.add_run("• ")
    run_bullet.font.size = Pt(font_size)
    run_bullet.font.color.rgb = TEAL
    run_bullet.font.name = "Calibri"
    
    if bold_prefix:
        run_b = p.add_run(bold_prefix)
        run_b.font.size = Pt(font_size)
        run_b.font.bold = True
        run_b.font.color.rgb = color
        run_b.font.name = "Calibri"
    
    run_text = p.add_run(text)
    run_text.font.size = Pt(font_size)
    run_text.font.color.rgb = color
    run_text.font.name = "Calibri"
    return p


def add_section_header(doc, text, color=TEAL):
    """Add a section header with teal color."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = Pt(14)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = color
    run.font.name = "Calibri"
    # Add a thin line below
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(2)
    p2.paragraph_format.space_before = Pt(0)
    # Use a border-bottom on the paragraph
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="6" w:space="1" w:color="0D9488"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    doc.element.body.remove(p2._p)
    return p


def create_summary():
    doc = Document()
    
    # ── Page Setup — narrow margins for one-page fit ──
    section = doc.sections[0]
    section.page_width = Inches(11.69)   # A4 landscape
    section.page_height = Inches(8.27)
    section.orientation = WD_ORIENT.LANDSCAPE
    section.top_margin = Cm(1.0)
    section.bottom_margin = Cm(0.8)
    section.left_margin = Cm(1.2)
    section.right_margin = Cm(1.2)
    
    # ── Default font ──
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(9)
    style.paragraph_format.space_after = Pt(2)
    style.paragraph_format.space_before = Pt(0)
    
    # ══════════════════════════════════════════════════════════
    # TITLE BAR
    # ══════════════════════════════════════════════════════════
    title_table = doc.add_table(rows=1, cols=2)
    title_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    title_table.columns[0].width = Inches(7.5)
    title_table.columns[1].width = Inches(3.0)
    
    # Set full-width dark background
    for cell in title_table.rows[0].cells:
        set_cell_bg(cell, NAVY_BG)
        cell.vertical_alignment = 1  # CENTER
    
    # Left: Title
    tc = title_table.cell(0, 0)
    tc.width = Inches(7.5)
    p = tc.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("INTELLI-CREDIT")
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x14, 0xB8, 0xA6)
    run.font.name = "Calibri"
    
    p2 = tc.add_paragraph()
    p2.paragraph_format.space_after = Pt(6)
    run2 = p2.add_run("AI-Powered Loan Default Risk Prediction Engine")
    run2.font.size = Pt(12)
    run2.font.color.rgb = WHITE
    run2.font.name = "Calibri"
    
    # Right: Hackathon info
    rc = title_table.cell(0, 1)
    rc.width = Inches(3.0)
    p = rc.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run("National Level AI Grand Challenge 2026")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
    run.font.name = "Calibri"
    p2 = rc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p2.paragraph_format.space_after = Pt(6)
    run2 = p2.add_run("FinTech-2: Loan Default Risk Prediction")
    run2.font.size = Pt(9)
    run2.font.bold = True
    run2.font.color.rgb = RGBColor(0x14, 0xB8, 0xA6)
    run2.font.name = "Calibri"
    
    # Small spacer
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)
    spacer.paragraph_format.space_before = Pt(2)
    
    # ══════════════════════════════════════════════════════════
    # MAIN CONTENT — 2-column layout using table
    # ══════════════════════════════════════════════════════════
    main_table = doc.add_table(rows=1, cols=2)
    main_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    main_table.columns[0].width = Inches(5.2)
    main_table.columns[1].width = Inches(5.0)
    
    left = main_table.cell(0, 0)
    right = main_table.cell(0, 1)
    left.width = Inches(5.2)
    right.width = Inches(5.0)
    
    # Remove default borders
    for cell in [left, right]:
        for edge in ['top', 'bottom', 'left', 'right']:
            set_cell_border(cell, **{edge: {"val": "none", "sz": "0", "color": "FFFFFF"}})
    
    # ──────────────────────────────────────────────────────────
    # LEFT COLUMN
    # ──────────────────────────────────────────────────────────
    
    # ── Problem ──
    p = left.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("THE PROBLEM")
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xEF, 0x44, 0x44)
    run.font.name = "Calibri"
    
    problems = [
        ("Manual process: ", "Credit officers read 8+ documents per application — takes days"),
        ("Hidden fraud: ", "Circular trading, shell companies, related party concealment — invisible in manual review"),
        ("No cross-check: ", "Company self-reports ₹247Cr revenue; GST portal shows ₹198Cr — nobody catches the gap"),
        ("No traceability: ", "When loans default, no audit trail of why it was approved"),
    ]
    for bold_p, text in problems:
        add_bullet(left, text, font_size=8, bold_prefix=bold_p, indent=0.15)
    
    # ── Proposed Solution ──
    p = left.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("PROPOSED SOLUTION")
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = TEAL
    run.font.name = "Calibri"
    
    solutions = [
        ("9 Parallel AI Workers ", "— read all 8 documents simultaneously, not sequentially"),
        ("4-Way Cross-Verification ", "— every revenue figure checked across Annual Report, ITR, GST, Bank Statement"),
        ("Knowledge Graph (Neo4j) ", "— maps companies, directors, suppliers into a graph; reasons over hidden connections"),
        ("5 Graph Reasoning Passes ", "— cascade risk, circular trading, concentration, temporal patterns, positive signals"),
        ("3 ML Models on Indian Data ", "— DOMINANT GNN (fraud graphs), Isolation Forest (ratio anomalies), FinBERT (buried risk text)"),
        ("5Cs Scoring (0–850) ", "— Capacity, Character, Capital, Collateral, Conditions + Compound — every point traceable"),
        ("Ticket Layer ", "— AI raises conflicts to human officer; resolutions stored as precedents"),
        ("Live AI Thinking ", "— officer watches reasoning in real-time via WebSocket chatbot feed"),
        ("Auto CAM Generation ", "— Credit Appraisal Memo in Indian banking format with citations"),
        ("Research Layer ", "— checks MCA21, SEBI, RBI, NJDG, GST portals + Tavily/Exa AI search"),
    ]
    for bold_p, text in solutions:
        add_bullet(left, text, font_size=8, bold_prefix=bold_p, color=SLATE_700, indent=0.15)
    
    # ── Key Differentiators ──
    p = left.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("KEY DIFFERENTIATORS")
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xA7, 0x8B, 0xFA)
    run.font.name = "Calibri"
    
    # Differentiator cards as a mini table
    diff_table = doc.add_table(rows=2, cols=3)
    # Move table into left cell
    left._tc.append(diff_table._tbl)
    
    diffs = [
        ("🔄 Parallel Agents", "9 Celery workers fire at once — one per document type. Consolidator merges & detects contradictions.", TEAL_BG),
        ("🧠 Live AI Thinking", "Every agent broadcasts reasoning via Redis → WebSocket. Officer sees: 'Found revenue gap: 20%'", BLUE_BG),
        ("🕸️ Graph Reasoning", "5 passes over Neo4j knowledge graph. Catches circular trading (A→B→C→A) invisible to document readers.", PURPLE_BG),
        ("🎯 5Cs Framework", "Indian banking standard. 850 points. Every single point traces to document + page + excerpt.", LIGHT_TEAL_BG),
        ("🎫 Ticket Layer", "AI knows when it's uncertain — raises tickets. Human resolves with evidence. Builds precedent store.", AMBER_BG),
        ("📄 Auto CAM", "Full Credit Appraisal Memo generated in Word/PDF. Every claim hyperlinked to source. Indian banking format.", GRAY_BG),
    ]
    
    for i, (title, desc, bg) in enumerate(diffs):
        row, col = divmod(i, 3)
        cell = diff_table.cell(row, col)
        set_cell_bg(cell, bg)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(title)
        run.font.size = Pt(7.5)
        run.font.bold = True
        run.font.color.rgb = DARK_NAVY
        run.font.name = "Calibri"
        p2 = cell.add_paragraph()
        p2.paragraph_format.space_after = Pt(3)
        run2 = p2.add_run(desc)
        run2.font.size = Pt(6.5)
        run2.font.color.rgb = SLATE_700
        run2.font.name = "Calibri"
    
    # ──────────────────────────────────────────────────────────
    # RIGHT COLUMN — Flowcharts + Tech + Impact
    # ──────────────────────────────────────────────────────────
    
    # ── Architecture Flowcharts ──
    p = right.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run("ARCHITECTURE FLOWCHARTS")
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = TEAL
    run.font.name = "Calibri"
    
    # Check for flowchart images
    base = r"c:\Users\harir\Downloads\ALL\projects\yuvaanai-hyd"
    flowchart_files = [
        ("flowchart_1.png", "3-Pillar Engine"),
        ("flowchart_2.png", "Parallel Execution Flow"),
        ("flowchart_3.png", "Scoring Pipeline"),
        ("flowchart_4.png", "Graph Reasoning + ML"),
    ]
    
    # Create 2x2 grid for flowcharts
    fc_table = doc.add_table(rows=2, cols=2)
    right._tc.append(fc_table._tbl)
    
    for i, (fname, label) in enumerate(flowchart_files):
        row, col = divmod(i, 2)
        cell = fc_table.cell(row, col)
        set_cell_bg(cell, GRAY_BG)
        
        fpath = os.path.join(base, fname)
        if os.path.exists(fpath):
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(3)
            run = p.add_run()
            run.add_picture(fpath, width=Inches(2.2))
        else:
            # Placeholder
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(15)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(f"[ INSERT: {label} ]")
            run.font.size = Pt(8)
            run.font.bold = True
            run.font.color.rgb = SLATE_500
            run.font.name = "Calibri"
            
            p2 = cell.add_paragraph()
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p2.paragraph_format.space_after = Pt(15)
            run2 = p2.add_run(f"Save as {fname}")
            run2.font.size = Pt(6.5)
            run2.font.italic = True
            run2.font.color.rgb = SLATE_500
            run2.font.name = "Calibri"
    
    # ── Tech Stack + Impact (compact) ──
    p = right.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("TECH STACK")
    run.font.size = Pt(9)
    run.font.bold = True
    run.font.color.rgb = TEAL
    run.font.name = "Calibri"
    
    p = right.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = Pt(10)
    
    tech_items = [
        ("AI/ML: ", "LangGraph + LangChain | PyTorch Geometric (GNN) | Isolation Forest | FinBERT | all-MiniLM-L6-v2"),
        ("Backend: ", "Python 3.12 | FastAPI | Celery 5 | Redis 7"),
        ("Storage: ", "Neo4j 5.14 | PostgreSQL 15 | ChromaDB 0.5 | Elasticsearch 8.12"),
        ("Frontend: ", "React 18 | TailwindCSS 3 | WebSocket | Recharts"),
        ("Infra: ", "Docker Compose (10 services, one-command deploy)"),
    ]
    for bold_p, text in tech_items:
        p = right.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.line_spacing = Pt(9)
        p.paragraph_format.left_indent = Inches(0.15)
        
        run = p.add_run(bold_p)
        run.font.size = Pt(6.5)
        run.font.bold = True
        run.font.color.rgb = DARK_NAVY
        run.font.name = "Calibri"
        run2 = p.add_run(text)
        run2.font.size = Pt(6.5)
        run2.font.color.rgb = SLATE_700
        run2.font.name = "Calibri"
    
    # ── Impact ──
    p = right.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("IMPACT & BENEFITS")
    run.font.size = Pt(9)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x10, 0xB9, 0x81)
    run.font.name = "Calibri"
    
    impacts = [
        ("Days → Minutes: ", "Parallel processing + AI eliminates manual bottleneck"),
        ("Fraud Detection: ", "Graph reasoning catches circular trading invisible to document readers"),
        ("Full Traceability: ", "Every score point → document → page → excerpt. Regulator-audit ready"),
        ("Human-in-the-Loop: ", "AI knows when uncertain, raises tickets. Officer always has final say"),
        ("Scalable: ", "Celery handles 10+ concurrent assessments. All open-source, Docker-deployable"),
    ]
    for bold_p, text in impacts:
        add_bullet(right, text, font_size=7.5, bold_prefix=bold_p, color=SLATE_700, indent=0.15)
    
    # ── Research Sources (compact) ──
    p = right.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("RESEARCH & REFERENCES")
    run.font.size = Pt(8)
    run.font.bold = True
    run.font.color.rgb = SLATE_500
    run.font.name = "Calibri"
    
    refs = [
        "DOMINANT GNN — Ding et al., SDM 2019 (Graph anomaly detection)",
        "Isolation Forest — Liu et al., ICDM 2008 (Tabular anomaly)",
        "FinBERT — Araci 2019 / ProsusAI (Financial text sentiment)",
        "5Cs Framework — RBI Master Circular on Loans & Advances",
        "Govt Sources: MCA21, SEBI, RBI, NJDG, GST Portal",
    ]
    for ref in refs:
        p = right.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.line_spacing = Pt(8.5)
        p.paragraph_format.left_indent = Inches(0.15)
        run = p.add_run(f"▸ {ref}")
        run.font.size = Pt(6)
        run.font.color.rgb = SLATE_500
        run.font.name = "Calibri"
    
    # ══════════════════════════════════════════════════════════
    # FOOTER
    # ══════════════════════════════════════════════════════════
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_p.paragraph_format.space_before = Pt(4)
    footer_p.paragraph_format.space_after = Pt(0)
    run = footer_p.add_run("Team devLOVEpers  |  ")
    run.font.size = Pt(7)
    run.font.color.rgb = SLATE_500
    run.font.name = "Calibri"
    run2 = footer_p.add_run("Intelli-Credit: AI-Powered Loan Default Risk Prediction  |  ")
    run2.font.size = Pt(7)
    run2.font.bold = True
    run2.font.color.rgb = TEAL
    run2.font.name = "Calibri"
    run3 = footer_p.add_run("National Level AI Grand Challenge 2026")
    run3.font.size = Pt(7)
    run3.font.color.rgb = SLATE_500
    run3.font.name = "Calibri"
    
    # ── Save ──
    output_path = os.path.join(base, "Intelli_Credit_Summary.docx")
    doc.save(output_path)
    print(f"✅ Summary saved to: {output_path}")
    return output_path


if __name__ == "__main__":
    create_summary()
